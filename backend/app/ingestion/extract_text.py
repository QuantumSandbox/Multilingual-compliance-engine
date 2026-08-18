from __future__ import annotations

import re
from pathlib import Path

from app.config import settings, UPLOAD_DIR
from app.models import Document, Chunk
from app.db import SessionLocal
from app.ingestion.language import detect_language
from app.ingestion.translate import translate_text
from app.ingestion.chunking import chunk_text
from app.extraction.embeddings import embed_texts
from app.extraction.rules import looks_like_obligation
import uuid


def _read_pdf(path: Path) -> list[dict]:
    """Return list of {page, text} using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF is required for PDF ingestion (pip install pymupdf)")
    pages = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc):
            pages.append({"page": i + 1, "text": page.get_text("text")})
    return pages


def _read_docx(path: Path) -> list[dict]:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX ingestion")
    document = docx.Document(str(path))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    # one "page" = the whole doc, paragraphs indexed
    return [{"page": 1, "text": "\n".join(paras)}]


def _read_txt(path: Path) -> list[dict]:
    return [{"page": 1, "text": path.read_text(encoding="utf-8", errors="ignore")}]


def ingest_file(file_path: str, meta: dict) -> dict:
    """Extract text, detect language, translate, chunk, embed. Returns summary."""
    p = Path(file_path)
    ext = p.suffix.lower().lstrip(".")
    if ext == "pdf":
        pages = _read_pdf(p)
    elif ext in ("docx", "doc"):
        pages = _read_docx(p)
    elif ext == "txt":
        pages = _read_txt(p)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    full_text = "\n".join(pg["text"] for pg in pages)
    lang = detect_language(full_text)
    original_lang = lang

    # Translate non-English to English for working text; keep original for citation.
    working_pages = []
    for pg in pages:
        if lang != "en":
            try:
                en = translate_text(pg["text"], src=lang, tgt="en")
            except Exception:
                en = pg["text"]
        else:
            en = pg["text"]
        working_pages.append({"page": pg["page"], "text": en, "original": pg["text"]})

    # Build flat chunk list with page/paragraph coordinates.
    chunks = []
    for pg in working_pages:
        for para_idx, (en_seg, orig_seg) in enumerate(
            zip(chunk_text(pg["text"]), chunk_text(pg["original"])), start=1
        ):
            if not en_seg.strip():
                continue
            chunks.append(
                {
                    "page": pg["page"],
                    "paragraph": para_idx,
                    "text": en_seg.strip(),
                    "original_text": orig_seg.strip() or en_seg.strip(),
                }
            )

    return {
        "language": lang,
        "original_language": original_lang,
        "chunks": chunks,
        "page_count": len(pages),
    }


def save_document_and_chunks(meta: dict, ingestion: dict, db: SessionLocal = None) -> Document:
    own = db is None
    if own:
        db = SessionLocal()
    try:
        doc = Document(
            doc_id=meta["doc_id"],
            title=meta["title"],
            source_type=meta.get("source_type", "circular"),
            source_dept=meta.get("source_dept", ""),
            language="en" if ingestion["language"] != "en" else "en",
            original_language=ingestion["original_language"],
            version=meta.get("version", "v1"),
            family_id=meta.get("family_id", meta["doc_id"]),
            filename=meta["filename"],
            file_path=meta["file_path"],
            file_type=meta.get("file_type", "pdf"),
            sensitive=meta.get("sensitive", False),
            retention_days=meta.get("retention_days", settings.default_retention_days),
            uploaded_by=meta.get("uploaded_by"),
        )
        db.add(doc)
        db.flush()

        texts = [c["text"] for c in ingestion["chunks"]]
        vectors = embed_texts(texts) if texts else []
        for i, c in enumerate(ingestion["chunks"]):
            db.add(
                Chunk(
                    doc_id=doc.doc_id,
                    chunk_index=i,
                    page=c["page"],
                    paragraph=c["paragraph"],
                    text=c["text"],
                    original_text=c["original_text"],
                    embedding=vectors[i] if i < len(vectors) else None,
                    is_obligation=looks_like_obligation(c["text"]),
                )
            )
        doc.chunk_count = len(ingestion["chunks"])
        db.commit()
        db.refresh(doc)
        return doc
    finally:
        if own:
            db.close()
